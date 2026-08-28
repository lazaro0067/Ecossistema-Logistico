from datetime import datetime, timedelta
import io
import re
import sqlite3
import xml.etree.ElementTree as ET
import zipfile
import pandas as pd
import streamlit as st

# Importação condicional do docx para evitar falhas no Streamlit Cloud
try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 1. Configuração Inicial da Página & Design System Sênior CSS
st.set_page_config(
    page_title="Gestão DPO & Distribuição - Grupo Lima",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
    <style>
        .main {
            background-color: #f4f6f9;
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        }
        
        /* Estilização da Sidebar com Fundo Azul Sênior */
        section[data-testid="stSidebar"] {
            background-color: #0d2149;
            color: #ffffff;
        }
        section[data-testid="stSidebar"] p, 
        section[data-testid="stSidebar"] span {
            color: #ffffff !important;
        }

        /* Correção de Contraste para os Botões da Sidebar */
        section[data-testid="stSidebar"] button {
            background-color: #ffffff !important;
            color: #0d2149 !important;
            font-weight: 600 !important;
            border: 1px solid #cbd5e1 !important;
            border-radius: 8px !important;
        }
        section[data-testid="stSidebar"] button:hover {
            background-color: #e2e8f0 !important;
            color: #000000 !important;
        }
        section[data-testid="stSidebar"] button[kind="primary"] {
            background-color: #3b82f6 !important;
            color: #ffffff !important;
            border: 1px solid #2563eb !important;
        }

        /* Cartões Estilo Sênior Corporativo */
        .senior-card {
            background-color: #ffffff;
            border: 1px solid #e2e8f0;
            border-radius: 10px;
            padding: 20px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.03);
            margin-bottom: 16px;
            transition: all 0.2s ease-in-out;
        }
        .senior-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1), 0 4px 6px -2px rgba(0, 0, 0, 0.05);
            border-color: #3b82f6;
        }
        
        /* Cards com Cores Vivas */
        .card-vibrante-coral {
            background: linear-gradient(135deg, #ff6b6b 0%, #ee5253 100%);
            color: white;
            border-radius: 12px;
            padding: 16px;
            box-shadow: 0 6px 12px rgba(238, 82, 83, 0.3);
            margin-bottom: 12px;
        }
        .card-vibrante-verde {
            background: linear-gradient(135deg, #1dd1a1 0%, #10ac84 100%);
            color: white;
            border-radius: 12px;
            padding: 16px;
            box-shadow: 0 6px 12px rgba(16, 172, 132, 0.3);
            margin-bottom: 12px;
        }
        .card-vibrante-azul {
            background: linear-gradient(135deg, #2e86de 0%, #0abde3 100%);
            color: white;
            border-radius: 12px;
            padding: 16px;
            box-shadow: 0 6px 12px rgba(10, 189, 227, 0.3);
            margin-bottom: 12px;
        }
        .card-vibrante-amarelo {
            background: linear-gradient(135deg, #feca57 0%, #ff9f43 100%);
            color: white;
            border-radius: 12px;
            padding: 16px;
            box-shadow: 0 6px 12px rgba(255, 159, 67, 0.3);
            margin-bottom: 12px;
        }
    </style>
""", unsafe_allow_html=True)

DEPARTAMENTOS_DISPONIVEIS = [
    "Visão Geral (Dashboard)",
    "Puxada",
    "Ressuprimento",
    "Vendas",
    "Armazém & Estoque",
    "Distribuição (Entrega)",
    "Frota & Manutenção",
    "Financeiro & OBZ",
    "Compras & Insumos",
    "Gente & SSMA",
    "Relatórios & Bases Globais",
]

OPERACOES_DISPONIVEIS = [
    "Lima Rio Verde",
    "Lima Barreiras",
    "Lima São Félix",
    "Bahia",
]


# 2. Inicialização do Banco de Dados SQLite e Migrações Seguras
def init_db():
    conn = sqlite3.connect("puxada_ambev.db")
    cursor = conn.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS base_01_11 (
        cod_clean TEXT PRIMARY KEY,
        descricao TEXT,
        fator_hl REAL,
        cx_pallet REAL
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS base_linear (
        cod_clean TEXT PRIMARY KEY,
        tipo TEXT,
        categoria TEXT,
        linear_vendas REAL,
        dt_atualizacao TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS base_estoque_02 (
        operacao TEXT,
        cod_clean TEXT,
        descricao TEXT,
        inicial REAL,
        entrada REAL,
        saida REAL,
        disponivel REAL,
        dt_atualizacao TEXT,
        PRIMARY KEY (operacao, cod_clean)
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS pedidos_marcados (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        data_puxada TEXT,
        cod_clean TEXT,
        descricao TEXT,
        cx_solicitadas REAL,
        cx_marcadas REAL,
        hl_marcado REAL,
        status_item TEXT,
        numero_pedido TEXT,
        dt_atualizacao TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS agendamentos_descarga (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        placa TEXT,
        slot TEXT,
        data_descarga TEXT,
        tipo_carga TEXT,
        observacao TEXT,
        dt_atualizacao TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS carretas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        placa TEXT UNIQUE NOT NULL,
        modelo TEXT,
        capacidade_hl REAL,
        status TEXT DEFAULT 'Disponível'
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS transportadoras_gestao (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        nome TEXT UNIQUE NOT NULL,
        cnpj TEXT,
        contato TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS fabricas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT UNIQUE NOT NULL,
        cidade TEXT,
        uf TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS motoristas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        nome TEXT NOT NULL,
        cnh TEXT,
        telefone TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS vinculos_pedidos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        numero_pedido TEXT,
        data_puxada TEXT,
        placa TEXT,
        fabrica TEXT,
        transportadora TEXT,
        motorista TEXT,
        notas_fiscais TEXT,
        dt_atualizacao TEXT
    )""")

    try:
        cursor.execute("ALTER TABLE vinculos_pedidos ADD COLUMN notas_fiscais TEXT")
    except Exception:
        pass

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS operacoes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT UNIQUE NOT NULL,
        cnpj TEXT, cidade TEXT, uf TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS usuarios (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT UNIQUE NOT NULL,
        senha TEXT NOT NULL,
        email TEXT, 
        cargo TEXT, 
        perfil TEXT CHECK(perfil IN ('Master', 'Operacional')),
        e_aprovador TEXT DEFAULT 'Não', 
        alcada_reais REAL DEFAULT 0.0,
        permissoes_operacoes TEXT DEFAULT 'TODAS',
        permissoes_deptos TEXT DEFAULT 'TODOS',
        status TEXT DEFAULT 'Ativo'
    )""")

    try:
        cursor.execute("ALTER TABLE usuarios ADD COLUMN status TEXT DEFAULT 'Ativo'")
    except Exception:
        pass

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS cadastro_trechos_frete (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        trecho TEXT UNIQUE NOT NULL,
        origem TEXT,
        destino TEXT,
        transportadora TEXT NOT NULL,
        valor_frete REAL NOT NULL,
        aprovador TEXT NOT NULL
    )""")

    try:
        cursor.execute("ALTER TABLE cadastro_trechos_frete ADD COLUMN origem TEXT")
        cursor.execute("ALTER TABLE cadastro_trechos_frete ADD COLUMN destino TEXT")
    except Exception:
        pass

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS padroes_dpo (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        modulo TEXT,
        subbloco TEXT,
        titulo_padrao TEXT,
        conteudo_padrao TEXT,
        sugestoes_ia TEXT,
        dt_atualizacao TEXT,
        UNIQUE(operacao, modulo, subbloco)
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS metas_doi (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT, cod_prod TEXT, doi_meta REAL DEFAULT 7.0,
        UNIQUE(operacao, cod_prod)
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS cotacoes_frete (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT, origem TEXT, destino TEXT, data_requisicao TEXT,
        data_frete TEXT, motivo TEXT, transportadora TEXT, valor_negociado REAL,
        centro_custo TEXT, solicitante TEXT, aprovador TEXT, observacao TEXT,
        status TEXT DEFAULT 'Pendente Aprovação',
        cte_anexado TEXT, notas_fiscais_anexadas TEXT
    )""")

    try:
        cursor.execute("ALTER TABLE cotacoes_frete ADD COLUMN cte_anexado TEXT")
        cursor.execute("ALTER TABLE cotacoes_frete ADD COLUMN notas_fiscais_anexadas TEXT")
    except Exception:
        pass

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS historico_curva_abc (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        mes_ano TEXT,
        cod_clean TEXT,
        descricao TEXT,
        total_qtde REAL,
        pct_acumulado REAL,
        classe_abc TEXT,
        dt_atualizacao TEXT,
        UNIQUE(operacao, mes_ano, cod_clean)
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS layout_armazem (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        nome_area TEXT,
        nome_arquivo TEXT,
        tipo_arquivo TEXT,
        dados_blob BLOB,
        dt_atualizacao TEXT
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS gestao_ressuprimento_diario (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        data_registro TEXT,
        mes_ano TEXT,
        cesta TEXT,
        volume_sellin_hl REAL DEFAULT 0.0,
        volume_real_hl REAL DEFAULT 0.0,
        dt_atualizacao TEXT,
        UNIQUE(operacao, data_registro, cesta)
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS metas_ressuprimento_mensal (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        ano INTEGER,
        mes INTEGER,
        mes_ano TEXT,
        cesta TEXT,
        meta_volume_hl REAL DEFAULT 0.0,
        dt_atualizacao TEXT,
        UNIQUE(operacao, ano, mes, cesta)
    )""")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS politica_estoque_base (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        data_registro TEXT,
        cod_clean TEXT,
        sku_original TEXT,
        tipo TEXT,
        categoria TEXT,
        estoque REAL,
        demanda REAL,
        doi_atual REAL,
        pe_min_dias REAL,
        pe_obj_dias REAL,
        pe_max_dias REAL,
        pe_min_hl REAL,
        pe_obj_hl REAL,
        pe_max_hl REAL,
        dt_atualizacao TEXT
    )""")

    try:
        cursor.execute("ALTER TABLE politica_estoque_base ADD COLUMN data_registro TEXT")
    except Exception:
        pass

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS financeiro_contas_pagar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operacao TEXT,
        pacote TEXT,
        nbz TEXT,
        departamento TEXT,
        data_vencimento TEXT,
        documento TEXT,
        fornecedor_id TEXT,
        nome_fornecedor TEXT,
        historico TEXT,
        conta_gerencial TEXT,
        vbz TEXT,
        comprometido REAL DEFAULT 0.0,
        realizado REAL DEFAULT 0.0,
        valor_col_n REAL DEFAULT 0.0,
        usuario TEXT,
        dt_atualizacao TEXT
    )""")

    try:
        cursor.execute("ALTER TABLE financeiro_contas_pagar ADD COLUMN valor_col_n REAL DEFAULT 0.0")
    except Exception:
        pass

    empresas = [
        ("Lima Rio Verde", "12.345.678/0001-90", "Rio Verde", "GO"),
        ("Lima Barreiras", "98.765.432/0001-10", "Barreiras", "BA"),
        ("Lima São Félix", "45.678.912/0001-33", "São Félix do Coribe", "BA"),
    ]
    for emp in empresas:
        cursor.execute(
            "INSERT OR IGNORE INTO operacoes (nome, cnpj, cidade, uf) VALUES (?,?,?,?)",
            emp,
        )

    cursor.execute("SELECT count(*) FROM usuarios WHERE nome = 'admin'")
    if cursor.fetchone()[0] == 0:
        cursor.execute("""
        INSERT INTO usuarios (nome, senha, email, cargo, perfil, e_aprovador, alcada_reais, permissoes_operacoes, permissoes_deptos, status)
        VALUES ('admin', 'admin123', 'admin@grupolima.com.br', 'Administrador Master', 'Master', 'Sim', 9999999.0, 'TODAS', 'TODOS', 'Ativo')
        """)

    conn.commit()
    conn.close()


init_db()


# 3. Funções Auxiliares de Leitura e Tratamento Ultra-Robustas
def read_word_file(file_obj):
    if not HAS_DOCX:
        return "A biblioteca 'python-docx' não está instalada no servidor."
    try:
        file_obj.seek(0)
        doc = docx.Document(file_obj)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        return f"Erro ao ler arquivo do Word: {e}"


def robust_read_file(file_obj):
    filename = str(file_obj.name).lower()

    if filename.endswith(".xls"):
        for h in [0, 1, 2, 3, 4, 5]:
            try:
                file_obj.seek(0)
                df = pd.read_excel(file_obj, header=h, engine="xlrd", dtype=str)
                if df is not None and not df.empty and len(df.columns) > 1:
                    df = df.dropna(how="all", axis=1)
                    return df
            except Exception:
                continue

    if filename.endswith(".xlsx") or filename.endswith(".xls"):
        for h in [0, 1, 2, 3, 4, 5]:
            try:
                file_obj.seek(0)
                df = pd.read_excel(file_obj, header=h, engine="openpyxl", dtype=str)
                if df is not None and not df.empty and len(df.columns) > 1:
                    df = df.dropna(how="all", axis=1)
                    return df
            except Exception:
                try:
                    file_obj.seek(0)
                    df = pd.read_excel(file_obj, header=h, dtype=str)
                    if df is not None and not df.empty and len(df.columns) > 1:
                        df = df.dropna(how="all", axis=1)
                        return df
                except Exception:
                    continue

    if filename.endswith(".docx") or filename.endswith(".doc"):
        text = read_word_file(file_obj)
        lines = [line.split("|") for line in text.split("\n") if line.strip()]
        return pd.DataFrame(lines) if lines else pd.DataFrame()

    for sep_char in [";", ",", "\t"]:
        try:
            file_obj.seek(0)
            df = pd.read_csv(
                file_obj,
                sep=sep_char,
                encoding="utf-8-sig",
                engine="python",
                on_bad_lines="skip",
                dtype=str,
            )
            if df is not None and len(df.columns) > 1:
                return df
        except Exception:
            try:
                file_obj.seek(0)
                df = pd.read_csv(
                    file_obj,
                    sep=sep_char,
                    encoding="latin1",
                    engine="python",
                    on_bad_lines="skip",
                    dtype=str,
                )
                if df is not None and len(df.columns) > 1:
                    return df
            except Exception:
                continue

    file_obj.seek(0)
    return pd.read_excel(file_obj, dtype=str)


def parse_br_float(val):
    if pd.isna(val):
        return 0.0
    if isinstance(val, (int, float)):
        return float(val)

    s = str(val).strip()
    if not s:
        return 0.0

    s = re.sub(r"[R\$\s]", "", s)

    if "." in s and "," in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s and "." not in s:
        s = s.replace(",", ".")
    elif "." in s:
        if s.count(".") > 1 or re.match(r"^\d{1,3}(\.\d{3})+$", s):
            s = s.replace(".", "")

    try:
        return float(s)
    except Exception:
        s_clean = re.sub(r"[^\d,\.-]", "", s)
        if "." in s_clean and "," in s_clean:
            s_clean = s_clean.replace(".", "").replace(",", ".")
        elif "," in s_clean:
            s_clean = s_clean.replace(",", ".")
        else:
            if s_clean.count(".") > 1 or re.match(
                r"^\d{1,3}(\.\d{3})+$", s_clean
            ):
                s_clean = s_clean.replace(".", "")
        try:
            return float(s_clean)
        except Exception:
            return 0.0


def formatar_br(val):
    try:
        if pd.isna(val):
            return "R$ 0,00"
        val_float = float(val)
        s_base = f"{val_float:,.2f}"
        s_base = s_base.replace(",", "X").replace(".", ",").replace("X", ".")
        return f"R$ {s_base}"
    except Exception:
        return str(val)


def formatar_inteiro_br(val):
    try:
        if pd.isna(val):
            return "0"
        val_float = float(val)
        s_base = f"{round(val_float):,}"
        s_base = s_base.replace(",", ".")
        return s_base
    except Exception:
        return str(val)


def classificar_tipo_sku(desc_or_sku):
    d = str(desc_or_sku).upper()
    if (
        "CERVEJA" in d
        or "CHOPP" in d
        or "BRAHMA" in d
        or "SKOL" in d
        or "BUDWEISER" in d
        or "SPATEN" in d
        or "CORONA" in d
        or "BECK" in d
        or "RET" in d
        or "KEG" in d
    ):
        return "CERVEJA"
    elif (
        "REFRIGERANTE" in d
        or "GUARANA" in d
        or "PEPSI" in d
        or "SUKITA" in d
        or "H2O" in d
        or "AGUA" in d
        or "ENERGETICO" in d
        or "TONICA" in d
        or "NAB" in d
    ):
        return "NAB"
    elif "MARKETPLACE" in d or "PIRACANJUBA" in d or "RED BULL" in d:
        return "MARKETPLACE"
    else:
        return "OUTROS"


def classificar_categoria_detalhada(desc_or_sku):
    d = str(desc_or_sku).upper()
    if "RET" in d or "RGB" in d:
        return "Retornável (Ret)"
    elif "DESC" in d or "LATA" in d or "LONG" in d or "PET" in d:
        return "Descartável (Desc)"
    elif "CERVEJA" in d or "CHOPP" in d or "BRAHMA" in d or "SKOL" in d or "SPATEN" in d or "BUD" in d:
        return "Cerveja"
    elif "NAB" in d or "REFRIGERANTE" in d or "PEPSI" in d or "GUARANA" in d or "AGUA" in d:
        return "NAB"
    else:
        return "Outros"


def extract_ambev_brand(desc):
    d = str(desc).upper().strip()
    if d.startswith("BC ") or "BRAHMA" in d or d.startswith("BR "):
        return "BRAHMA"
    elif d.startswith("SK ") or "SKOL" in d:
        return "SKOL"
    elif d.startswith("ANT ") or "ANTARCTICA" in d:
        return "ANTARCTICA"
    elif d.startswith("ORG ") or "ORIGINAL" in d:
        return "ORIGINAL"
    elif d.startswith("BUD") or "BUDWEISER" in d:
        return "BUDWEISER"
    elif d.startswith("SPAT") or "SPATEN" in d:
        return "SPATEN"
    elif d.startswith("COR ") or "CORONA" in d:
        return "CORONA"
    elif "ARTOIS" in d or d.startswith("SLA ") or "STELLA" in d:
        return "STELLA ARTOIS"
    elif d.startswith("SU ") or d.startswith("SUK") or "SUKITA" in d:
        return "SUKITA"
    elif d.startswith("PC ") or "PEPSI" in d:
        return "PEPSI"
    elif d.startswith("GCA ") or "GUARANA" in d:
        return "GUARANÁ"
    elif d.startswith("CHP BR") or "CHOPP" in d:
        return "CHOPP BRAHMA"
    elif d.startswith("BECK") or "BECKS" in d:
        return "BECKS"
    elif "H2O" in d:
        return "H2OH!"
    else:
        return "OUTROS"


def highlight_curva_abc(val):
    if val == "A":
        return "background-color: #d4edda; color: #155724; font-weight: bold;"
    elif val == "B":
        return "background-color: #fff3cd; color: #856404; font-weight: bold;"
    elif val == "C":
        return "background-color: #f8d7da; color: #721c24; font-weight: bold;"
    return ""


def aplicar_estilo_tabela(df_styled, subset_cols):
    try:
        return df_styled.map(highlight_curva_abc, subset=subset_cols)
    except AttributeError:
        return df_styled.applymap(highlight_curva_abc, subset=subset_cols)


def gerar_excel(df):
    output = io.BytesIO()
    try:
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Dados")
    except Exception:
        df.to_csv(output, index=False, sep=";")
    return output.getvalue()


def gerar_csv(df):
    return df.to_csv(index=False, sep=";", encoding="utf-8-sig").encode("utf-8-sig")


def render_botoes_download(df_export, nome_base):
    c_dl1, c_dl2 = st.columns(2)
    c_dl1.download_button(
        f"📥 Baixar {nome_base} (.xlsx)",
        data=gerar_excel(df_export),
        file_name=f"{nome_base}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )
    c_dl2.download_button(
        f"📥 Baixar {nome_base} (.csv)",
        data=gerar_csv(df_export),
        file_name=f"{nome_base}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
        mime="text/csv",
        use_container_width=True,
    )


def salvar_base_01_11(f_01):
    df_01 = robust_read_file(f_01)
    col_cod = [
        c for c in df_01.columns if "Código" in str(c) or "Cod" in str(c)
    ][0]
    df_01["cod_clean"] = df_01[col_cod].astype(str).str.strip()

    col_q = df_01.columns[16] if len(df_01.columns) > 16 else df_01.columns[-1]
    df_01["fator_hl"] = (
        df_01[col_q].astype(str).str.replace(",", ".").astype(float)
    )

    col_cx = [
        c
        for c in df_01.columns
        if "Caixas Pallet" in str(c) or "Pallet" in str(c)
    ]
    if col_cx:
        df_01["cx_pallet"] = pd.to_numeric(
            df_01[col_cx[0]], errors="coerce"
        ).fillna(1)
    else:
        df_01["cx_pallet"] = 1.0

    df_01["cx_pallet"] = df_01["cx_pallet"].apply(lambda x: x if x > 0 else 1)

    col_desc = [
        c for c in df_01.columns if "Descrição" in str(c) or "Desc" in str(c)
    ][0]
    df_sub = df_01[["cod_clean", col_desc, "fator_hl", "cx_pallet"]].dropna(
        subset=["cod_clean"]
    )

    conn = sqlite3.connect("puxada_ambev.db")
    cursor = conn.cursor()
    for _, r in df_sub.iterrows():
        cursor.execute(
            """
        INSERT INTO base_01_11 (cod_clean, descricao, fator_hl, cx_pallet)
        VALUES (?, ?, ?, ?)
        ON CONFLICT(cod_clean) DO UPDATE SET
            descricao=excluded.descricao, fator_hl=excluded.fator_hl, cx_pallet=excluded.cx_pallet
        """,
            (
                str(r["cod_clean"]),
                str(r[col_desc]),
                float(r["fator_hl"]),
                float(r["cx_pallet"]),
            ),
        )
    conn.commit()
    conn.close()


def salvar_base_linear(f_lin):
    df_lin = robust_read_file(f_lin)

    cols_str = [str(c).strip().lower() for c in df_lin.columns]
    col_cod = None
    for i, c in enumerate(cols_str):
        if "cód" in c or "cod" in c or "item" in c or "produto" in c:
            col_cod = df_lin.columns[i]
            break
    if col_cod is None:
        col_cod = df_lin.columns[0]

    if len(df_lin.columns) > 4:
        col_vendas = df_lin.columns[4]
    else:
        col_vendas = (
            df_lin.columns[1] if len(df_lin.columns) > 1 else df_lin.columns[0]
        )

    df_lin["cod_clean"] = df_lin[col_cod].astype(str).str.strip()
    df_lin["linear_vendas"] = df_lin[col_vendas].apply(parse_br_float)
    df_lin["Tipo"] = df_lin.get("Tipo", df_lin.get("tipo", pd.Series())).fillna(
        "OUTROS"
    )
    df_lin["Categoria"] = df_lin.get(
        "Categoria", df_lin.get("categoria", pd.Series())
    ).fillna("OUTROS")

    dt_now = datetime.now().strftime("%d/%m/%Y %H:%M")

    conn = sqlite3.connect("puxada_ambev.db")
    cursor = conn.cursor()
    for _, r in df_lin.dropna(subset=["cod_clean"]).iterrows():
        cursor.execute(
            """
        INSERT INTO base_linear (cod_clean, tipo, categoria, linear_vendas, dt_atualizacao)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(cod_clean) DO UPDATE SET
            tipo=excluded.tipo, categoria=excluded.categoria,
            linear_vendas=excluded.linear_vendas, dt_atualizacao=excluded.dt_atualizacao
        """,
            (
                str(r["cod_clean"]),
                str(r.get("Tipo", "OUTROS")),
                str(r.get("Categoria", "OUTROS")),
                float(r["linear_vendas"]),
                dt_now,
            ),
        )
    conn.commit()
    conn.close()


def salvar_base_estoque_02(f_02, operacao):
    df_02 = robust_read_file(f_02)

    col_cod = [
        c for c in df_02.columns if "Cod" in str(c) or "COD" in str(c)
    ][0]
    col_desc = [c for c in df_02.columns if "Desc" in str(c)][0]
    col_init = [c for c in df_02.columns if "Inic" in str(c)][0]
    col_ent = [c for c in df_02.columns if "Ent" in str(c)][0]
    col_sai = [
        c for c in df_02.columns if "Saida" in str(c) or "Sai" in str(c)
    ][0]
    col_disp = [c for c in df_02.columns if "Disp" in str(c)][0]

    df_02["cod_clean"] = df_02[col_cod].astype(str).str.strip()
    df_02["Inicial"] = pd.to_numeric(df_02[col_init], errors="coerce").fillna(0)
    df_02["Ent."] = pd.to_numeric(df_02[col_ent], errors="coerce").fillna(0)
    df_02["Saidas"] = pd.to_numeric(df_02[col_sai], errors="coerce").fillna(0)
    df_02["Disp."] = pd.to_numeric(df_02[col_disp], errors="coerce").fillna(0)

    dt_now = datetime.now().strftime("%d/%m/%Y %H:%M")

    conn = sqlite3.connect("puxada_ambev.db")
    cursor = conn.cursor()
    for _, r in df_02.dropna(subset=["cod_clean"]).iterrows():
        cursor.execute(
            """
        INSERT INTO base_estoque_02 (operacao, cod_clean, descricao, inicial, entrada, saida, disponivel, dt_atualizacao)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(operacao, cod_clean) DO UPDATE SET
            descricao=excluded.descricao, inicial=excluded.inicial, entrada=excluded.entrada,
            saida=excluded.saida, disponivel=excluded.disponivel, dt_atualizacao=excluded.dt_atualizacao
        """,
            (
                operacao,
                str(r["cod_clean"]),
                str(r[col_desc]),
                float(r["Inicial"]),
                float(r["Ent."]),
                float(r["Saidas"]),
                float(r["Disp."]),
                dt_now,
            ),
        )
    conn.commit()
    conn.close()


def salvar_pedidos_marcados(f_pedidos, operacao):
    df_pedidos = robust_read_file(f_pedidos)

    col_cod_r = (
        df_pedidos.columns[17]
        if len(df_pedidos.columns) > 17
        else [
            c
            for c in df_pedidos.columns
            if "Código" in str(c) or "Cod" in str(c)
        ][0]
    )
    col_marc_w = (
        df_pedidos.columns[22]
        if len(df_pedidos.columns) > 22
        else [c for c in df_pedidos.columns if "Marcado" in str(c)][0]
    )
    col_desc = [
        c
        for c in df_pedidos.columns
        if "Produto" in str(c) or "Desc" in str(c)
    ][0]
    col_dt = [
        c
        for c in df_pedidos.columns
        if "Data Puxada" in str(c) or "Data" in str(c)
    ][0]
    col_solic = [
        c
        for c in df_pedidos.columns
        if "QtdeSKUs - Item" in str(c) or "Solicitado" in str(c)
    ][0]
    col_hl = [
        c for c in df_pedidos.columns if "HL" in str(c) or "Hecto" in str(c)
    ][0]

    conn = sqlite3.connect("puxada_ambev.db")
    cursor = conn.cursor()

    cursor.execute("DELETE FROM pedidos_marcados WHERE operacao=?", (operacao,))

    dt_now = datetime.now().strftime("%d/%m/%Y %H:%M")
    for _, r in df_pedidos.iterrows():
        cod_raw = str(r.get(col_cod_r, "")).strip()
        if not cod_raw:
            continue

        cx_marcadas_w = parse_br_float(r.get(col_marc_w))

        cursor.execute(
            """
        INSERT INTO pedidos_marcados (operacao, data_puxada, cod_clean, descricao, cx_solicitadas, cx_marcadas, hl_marcado, status_item, numero_pedido, dt_atualizacao)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
            (
                operacao,
                str(r.get(col_dt)).strip(),
                cod_raw,
                str(r.get(col_desc)).strip(),
                parse_br_float(r.get(col_solic)),
                cx_marcadas_w,
                parse_br_float(r.get(col_hl)),
                str(r.get("Status - Item", "")).strip(),
                str(r.get("Nº - Pedido", "")).strip(),
                dt_now,
            ),
        )

    conn.commit()
    conn.close()


def processar_politica_estoque_upload(f_pol, operacao):
    df_raw = robust_read_file(f_pol)
    dt_now = datetime.now().strftime("%d/%m/%Y %H:%M")
    conn = sqlite3.connect("puxada_ambev.db")
    cursor = conn.cursor()

    count = 0
    for _, r in df_raw.iterrows():
        try:
            raw_date = r.iloc[0]
            if pd.notna(raw_date):
                dt_parsed = pd.to_datetime(raw_date, errors="coerce")
                date_str = (
                    dt_parsed.strftime("%Y-%m-%d")
                    if pd.notna(dt_parsed)
                    else str(raw_date)[:10]
                )
            else:
                date_str = datetime.now().strftime("%Y-%m-%d")

            sku_str = str(r.iloc[2]).strip()
            partes = sku_str.split("/")
            if len(partes) >= 2:
                sub_part = partes[1].strip()
                cod_match = re.search(r"^\d+", sub_part)
                cod_clean = cod_match.group() if cod_match else "0"
            else:
                cod_match = re.search(r"^\d+", sku_str)
                cod_clean = cod_match.group() if cod_match else "0"

            if cod_clean == "0":
                continue

            tipo_classificado = classificar_tipo_sku(sku_str)
            cat = str(r.iloc[3]).strip() if pd.notna(r.iloc[3]) else "GERAL"
            est = parse_br_float(r.iloc[5])
            dem = parse_br_float(r.iloc[6])
            doi = parse_br_float(r.iloc[7])
            pe_min_d = parse_br_float(r.iloc[8])
            pe_obj_d = parse_br_float(r.iloc[9])
            pe_max_d = parse_br_float(r.iloc[10])
            pe_min_h = parse_br_float(r.iloc[11])
            pe_obj_h = parse_br_float(r.iloc[12])
            pe_max_h = parse_br_float(r.iloc[13])

            cursor.execute(
                """
            INSERT INTO politica_estoque_base (operacao, data_registro, cod_clean, sku_original, tipo, categoria, estoque, demanda, doi_atual, pe_min_dias, pe_obj_dias, pe_max_dias, pe_min_hl, pe_obj_hl, pe_max_hl, dt_atualizacao)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
                (
                    operacao,
                    date_str,
                    cod_clean,
                    sku_str,
                    tipo_classificado,
                    cat,
                    est,
                    dem,
                    doi,
                    pe_min_d,
                    pe_obj_d,
                    pe_max_d,
                    pe_min_h,
                    pe_obj_h,
                    pe_max_h,
                    dt_now,
                ),
            )
            count += 1
        except Exception:
            continue

    conn.commit()
    conn.close()
    return count


def processar_financeiro_upload(f_fin, operacao):
    df_raw = robust_read_file(f_fin)
    dt_now = datetime.now().strftime("%d/%m/%Y %H:%M")
    conn = sqlite3.connect("puxada_ambev.db")
    cursor = conn.cursor()

    cursor.execute("DELETE FROM financeiro_contas_pagar WHERE operacao = ?", (operacao,))

    count = 0
    for _, r in df_raw.iterrows():
        try:
            pacote = str(r.get("Pacote", r.iloc[0] if len(r) > 0 else "")).strip()
            nome_forn = str(r.iloc[1] if len(r) > 1 else r.get("Fornecedor", "")).strip()
            nbz = str(r.get("NBZ", r.iloc[2] if len(r) > 2 else "")).strip()
            dept = str(r.get("Departamento", r.iloc[3] if len(r) > 3 else "")).strip()
            
            raw_dt = r.iloc[10] if len(r) > 10 else r.get("Data", datetime.now())
            dt_parsed = pd.to_datetime(raw_dt, errors="coerce")
            data_venc = dt_parsed.strftime("%Y-%m-%d") if pd.notna(dt_parsed) else datetime.now().strftime("%Y-%m-%d")

            doc = str(r.get("Documento", r.iloc[4] if len(r) > 4 else "")).strip()
            forn_id = str(r.get("Fornecedor ID", r.iloc[5] if len(r) > 5 else "")).strip()
            hist = str(r.get("Historico", r.iloc[7] if len(r) > 7 else "")).strip()
            conta_ger = str(r.get("Conta Gerencial", r.iloc[8] if len(r) > 8 else "")).strip()
            vbz = str(r.get("VBZ", r.iloc[9] if len(r) > 9 else "")).strip()
            
            comprometido = parse_br_float(r.get("Comprometido", r.iloc[10] if len(r) > 10 else 0))
            realizado = parse_br_float(r.get("Realizado", r.iloc[11] if len(r) > 11 else 0))
            val_n = parse_br_float(r.iloc[13]) if len(r) > 13 else comprometido
            usuario = str(r.get("Usuario", r.iloc[12] if len(r) > 12 else "Sistema")).strip()

            cursor.execute(
                """
                INSERT INTO financeiro_contas_pagar (operacao, pacote, nbz, departamento, data_vencimento, documento, fornecedor_id, nome_fornecedor, historico, conta_gerencial, vbz, comprometido, realizado, valor_col_n, usuario, dt_atualizacao)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (operacao, pacote, nbz, dept, data_venc, doc, forn_id, nome_forn, hist, conta_ger, vbz, comprometido, realizado, val_n, usuario, dt_now)
            )
            count += 1
        except Exception:
            continue

    conn.commit()
    conn.close()
    return count


def carregar_estoque_consolidado(operacao):
    conn = sqlite3.connect("puxada_ambev.db")

    ops_filtro = [operacao]
    if operacao == "Bahia":
        ops_filtro = [
            "Lima Barreiras",
            "Lima São Félix",
            "Lima Bahia",
            "Lima Bahia Samavi",
        ]

    placeholders = ",".join(["?"] * len(ops_filtro))
    query = f"""
    SELECT 
        e.cod_clean AS cod_clean,
        COALESCE(e.descricao, b_01.descricao, 'PRODUTO') AS descricao,
        COALESCE(l.tipo, 'OUTROS') AS tipo,
        COALESCE(l.categoria, 'OUTROS') AS categoria,
        CAST(COALESCE(e.inicial, 0) AS INTEGER) AS inicial,
        CAST(COALESCE(e.entrada, 0) AS INTEGER) AS entrada,
        CAST(COALESCE(e.saida, 0) AS INTEGER) AS saida,
        CAST(COALESCE(e.disponivel, 0) AS INTEGER) AS disp,
        CAST(COALESCE(l.linear_vendas, 0) AS INTEGER) AS linear_vendas,
        COALESCE(b_01.fator_hl, 0.0) AS fator_hl,
        COALESCE(b_01.cx_pallet, 1.0) AS cx_pallet,
        e.dt_atualizacao AS dt_atualizacao
    FROM base_estoque_02 e
    LEFT JOIN base_01_11 b_01 ON e.cod_clean = b_01.cod_clean
    LEFT JOIN base_linear l ON e.cod_clean = l.cod_clean
    WHERE e.operacao IN ({placeholders})
    """

    df = pd.read_sql_query(query, conn, params=ops_filtro)
    
    query_marc = f"""
    SELECT cod_clean, data_puxada, SUM(cx_marcadas) as cx_marcadas
    FROM pedidos_marcados
    WHERE operacao IN ({placeholders})
    GROUP BY cod_clean, data_puxada
    """
    df_marc = pd.read_sql_query(query_marc, conn, params=ops_filtro)
    conn.close()

    if df.empty:
        return None

    df.columns = [str(c).lower() for c in df.columns]

    df["d0"] = 0.0
    df["d1"] = 0.0
    df["d2"] = 0.0

    if not df_marc.empty and "data_puxada" in df_marc.columns:
        datas_pux = sorted(df_marc["data_puxada"].dropna().unique())
        if len(datas_pux) > 0:
            d0_date = datas_pux[0]
            df_d0 = df_marc[df_marc["data_puxada"] == d0_date].groupby("cod_clean")["cx_marcadas"].sum().reset_index()
            df = pd.merge(df, df_d0.rename(columns={"cx_marcadas": "d0_val"}), on="cod_clean", how="left")
            if "d0_val" in df.columns:
                df["d0"] = df["d0_val"].fillna(0)
                df = df.drop(columns=["d0_val"])
        if len(datas_pux) > 1:
            d1_date = datas_pux[1]
            df_d1 = df_marc[df_marc["data_puxada"] == d1_date].groupby("cod_clean")["cx_marcadas"].sum().reset_index()
            df = pd.merge(df, df_d1.rename(columns={"cx_marcadas": "d1_val"}), on="cod_clean", how="left")
            if "d1_val" in df.columns:
                df["d1"] = df["d1_val"].fillna(0)
                df = df.drop(columns=["d1_val"])
        if len(datas_pux) > 2:
            d2_date = datas_pux[2]
            df_d2 = df_marc[df_marc["data_puxada"] == d2_date].groupby("cod_clean")["cx_marcadas"].sum().reset_index()
            df = pd.merge(df, df_d2.rename(columns={"cx_marcadas": "d2_val"}), on="cod_clean", how="left")
            if "d2_val" in df.columns:
                df["d2"] = df["d2_val"].fillna(0)
                df = df.drop(columns=["d2_val"])

    df["classe_abc"] = "C"
    df["total_puxada"] = df["d0"] + df["d1"] + df["d2"]
    df["estoque_projetado"] = df["disp"] + df["total_puxada"]

    df["doi_atual"] = df.apply(
        lambda r: round(r["disp"] / r["linear_vendas"], 1)
        if r["linear_vendas"] > 0
        else (999.0 if r["disp"] > 0 else 0.0),
        axis=1,
    )

    df["estoque_hl"] = (df["fator_hl"] * df["disp"]).round(2)
    df["marca"] = df["descricao"].apply(extract_ambev_brand)
    df["categoria_detalhada"] = df["descricao"].apply(classificar_categoria_detalhada)
    df["paletes_ocupados"] = (df["disp"] / df["cx_pallet"]).round(1)

    return df


def calcular_saude_estoque_dpo(df):
    if df is None or df.empty:
        return 0.0, 0, 0, 0
    total_skus = len(df)
    saudaveis = len(df[(df["doi_atual"] >= 3.0) & (df["doi_atual"] <= 15.0)])
    ruptura = len(df[df["doi_atual"] < 3.0])
    excesso = len(df[df["doi_atual"] > 15.0])
    pct_saude = (saudaveis / total_skus) * 100.0 if total_skus > 0 else 0.0
    return round(pct_saude, 1), saudaveis, ruptura, excesso


def gerar_analise_ia_financeiro(df_fin):
    if df_fin is None or df_fin.empty:
        return "Sem dados financeiros suficientes para análise da IA."
    
    tot_comp = df_fin["comprometido"].sum()
    tot_real = df_fin["realizado"].sum()
    pendente = tot_comp - tot_real
    
    maior_forn = df_fin.groupby("nome_fornecedor")["comprometido"].sum().idxmax() if not df_fin.empty else "N/A"
    maior_conta = df_fin.groupby("conta_gerencial")["comprometido"].sum().idxmax() if not df_fin.empty else "N/A"

    analise = [
        f"🤖 **Diagnóstico de Saúde Financeira e Contas a Pagar (IA)**:",
        f"- **Comprometido Total**: R$ {formatar_br(tot_comp)} | **Realizado**: R$ {formatar_br(tot_real)} | **Pendente**: R$ {formatar_br(pendente)}",
        f"- **Maior Fornecedor Comprometido**: {maior_forn}",
        f"- **Conta Gerencial de Maior Impacto**: {maior_conta}",
        f"- **Saúde de Caixa**: Unidade apresenta compromissos operacionais sob controle, recomenda-se monitorar os vencimentos diários para garantir capital de giro adequado."
    ]
    return "\n".join(analise)


# 5. Navegação por Pilhas de Histórico (Botão Voltar)
if "nav_stack" not in st.session_state:
    st.session_state["nav_stack"] = ["Visão Geral (Dashboard)"]


def navigate_to(page_name):
    if st.session_state["nav_stack"][-1] != page_name:
        st.session_state["nav_stack"].append(page_name)


def go_back():
    if len(st.session_state["nav_stack"]) > 1:
        st.session_state["nav_stack"].pop()


modo_comercial = False
modo_visualizacao_estatica = False
op_estatica = None

if "modo" in st.query_params:
    val_modo = st.query_params["modo"]
    if isinstance(val_modo, list):
        modo_comercial = "comercial" in val_modo
    else:
        modo_comercial = val_modo == "comercial"


def render_estoque_dia(unidade):
    st.subheader("Portal do RN - Consulta Comercial de Vendas (Com Marcações D0, D1 e D2)")

    df = carregar_estoque_consolidado(unidade)

    if df is not None and not df.empty:
        st.caption(f"🕒 **Última Atualização:** {df['dt_atualizacao'].iloc[0]}")

        def get_status(doi, disp):
            if disp == 0:
                return "🔴 Stock Out (Zerado)"
            elif doi < 3.0:
                return "🟡 Stock Low (Baixo)"
            elif doi <= 15.0:
                return "🟢 Stock Ideal"
            else:
                return "🔵 Stock Over (Excesso)"

        df["status"] = df.apply(
            lambda r: get_status(r["doi_atual"], r["disp"]), axis=1
        )

        total_skus_geral = len(df)
        stock_out_cnt = len(df[df["status"] == "🔴 Stock Out (Zerado)"])
        stock_low_cnt = len(df[df["status"] == "🟡 Stock Low (Baixo)"])
        stock_ideal_cnt = len(df[df["status"] == "🟢 Stock Ideal"])
        stock_over_cnt = len(df[df["status"] == "🔵 Stock Over (Excesso)"])

        pct_out = (stock_out_cnt / total_skus_geral * 100) if total_skus_geral > 0 else 0
        pct_low = (stock_low_cnt / total_skus_geral * 100) if total_skus_geral > 0 else 0
        pct_ideal = (stock_ideal_cnt / total_skus_geral * 100) if total_skus_geral > 0 else 0
        pct_over = (stock_over_cnt / total_skus_geral * 100) if total_skus_geral > 0 else 0

        if "rn_filtro_status" not in st.session_state:
            st.session_state["rn_filtro_status"] = "TODOS"

        k1, k2, k3, k4 = st.columns(4)

        with k1:
            if st.button(f"🔴 Stock Out\n{stock_out_cnt} SKUs ({pct_out:.1f}%)", use_container_width=True):
                st.session_state["rn_filtro_status"] = "🔴 Stock Out (Zerado)"
            st.markdown(f'<div class="card-vibrante-coral" style="text-align:center; margin-top:-10px;"><b>Crítico / Zerado ({pct_out:.1f}%)</b></div>', unsafe_allow_html=True)

        with k2:
            if st.button(f"🟡 Stock Low\n{stock_low_cnt} SKUs ({pct_low:.1f}%)", use_container_width=True):
                st.session_state["rn_filtro_status"] = "🟡 Stock Low (Baixo)"
            st.markdown(f'<div class="card-vibrante-amarelo" style="text-align:center; margin-top:-10px;"><b>Atenção / Baixo ({pct_low:.1f}%)</b></div>', unsafe_allow_html=True)

        with k3:
            if st.button(f"🟢 Stock Ideal\n{stock_ideal_cnt} SKUs ({pct_ideal:.1f}%)", use_container_width=True):
                st.session_state["rn_filtro_status"] = "🟢 Stock Ideal"
            st.markdown(f'<div class="card-vibrante-verde" style="text-align:center; margin-top:-10px;"><b>Saudável / Ideal ({pct_ideal:.1f}%)</b></div>', unsafe_allow_html=True)

        with k4:
            if st.button(f"🔵 Stock Over\n{stock_over_cnt} SKUs ({pct_over:.1f}%)", use_container_width=True):
                st.session_state["rn_filtro_status"] = "🔵 Stock Over (Excesso)"
            st.markdown(f'<div class="card-vibrante-azul" style="text-align:center; margin-top:-10px;"><b>Excesso / Overstock ({pct_over:.1f}%)</b></div>', unsafe_allow_html=True)

        st.divider()

        c_f1, c_f2, c_f3 = st.columns([2, 1, 1])
        busca = c_f1.text_input("🔍 Pesquisar por Código Completo ou Nome do Produto:")
        
        tipo_filtro_sel = c_f2.selectbox(
            "Tipo (Cerveja / NAB):", ["TODOS", "CERVEJA", "NAB"]
        )
        cat_det_sel = c_f3.selectbox(
            "Categoria (Ret / Desc):", ["TODAS", "Retornável (Ret)", "Descartável (Desc)", "Cerveja", "NAB"]
        )

        df_filtrado = df.copy()
        if st.session_state["rn_filtro_status"] != "TODOS":
            df_filtrado = df_filtrado[
                df_filtrado["status"] == st.session_state["rn_filtro_status"]
            ]

        if tipo_filtro_sel != "TODOS":
            df_filtrado = df_filtrado[df_filtrado["tipo"] == tipo_filtro_sel]

        if cat_det_sel != "TODAS":
            df_filtrado = df_filtrado[df_filtrado["categoria_detalhada"] == cat_det_sel]

        if busca:
            df_filtrado = df_filtrado[
                df_filtrado["cod_clean"].astype(str).str.contains(busca)
                | df_filtrado["descricao"].str.contains(busca, case=False)
            ]

        for _, r in df_filtrado.iterrows():
            if "Stock Out" in r["status"]:
                b_color = "#ee5253"
            elif "Stock Low" in r["status"]:
                b_color = "#ff9f43"
            elif "Stock Ideal" in r["status"]:
                b_color = "#10ac84"
            else:
                b_color = "#0abde3"

            disp_fmt = formatar_br(r["disp"])
            d0_fmt = formatar_br(r["d0"])
            d1_fmt = formatar_br(r["d1"])
            d2_fmt = formatar_br(r["d2"])
            proj_fmt = formatar_br(r["estoque_projetado"])

            st.markdown(
                f"""
                <div class="senior-card" style="border-left: 6px solid {b_color};">
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <span style="font-size: 11px; color: #444; font-weight: bold;">CÓD COMPLETO: {r['cod_clean']} | TIPO: {r['tipo']} | CAT: {r['categoria_detalhada']}</span>
                        <span style="font-size: 11px; font-weight: bold; background-color: #ffffff; color: #222; padding: 4px 10px; border-radius: 12px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">{r['status']}</span>
                    </div>
                    <div style="font-size: 16px; font-weight: bold; color: #1e293b; margin: 8px 0;">{r['descricao']}</div>
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-top: 10px; background-color: #f8fafc; padding: 10px; border-radius: 8px;">
                        <div><b>Estoque Físico:</b> <span style="color: #0288d1; font-size: 14px; font-weight: bold;">{disp_fmt} cx</span></div>
                        <div><b>D0:</b> <span style="color: #10ac84; font-size: 14px; font-weight: bold;">{d0_fmt} cx</span></div>
                        <div><b>D1:</b> <span style="color: #3b82f6; font-size: 14px; font-weight: bold;">{d1_fmt} cx</span></div>
                        <div><b>D2:</b> <span style="color: #8b5cf6; font-size: 14px; font-weight: bold;">{d2_fmt} cx</span></div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        render_botoes_download(df_filtrado, f"Estoque_Comercial_{unidade}")


if modo_comercial:
    st.title("Grupo Lima - Portal Comercial")
    unidade = st.selectbox(
        "Selecione a Unidade Operacional:", OPERACOES_DISPONIVEIS
    )
    render_estoque_dia(unidade)
    st.stop()


# 7. Autenticação e Navegação do Sistema
if "logado" not in st.session_state:
    st.session_state["logado"] = False

if not st.session_state["logado"]:
    st.markdown("""
        <div style="text-align: center; padding: 20px;">
            <h1 style="color: #0d2149;">Sistema Revenda - Grupo Lima</h1>
            <p style="color: #555;">Faça login com suas credenciais corporativas para acessar o sistema.</p>
        </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1.2, 1])
    with col2:
        with st.container():
            usuario = st.text_input("Usuário Corporativo", key="input_user_login")
            senha = st.text_input("Senha", type="password", key="input_pass_login")

            if st.button("🚀 Entrar no Sistema", use_container_width=True, key="btn_login_submit"):
                conn = sqlite3.connect("puxada_ambev.db")
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT id, nome, perfil, permissoes_operacoes, permissoes_deptos, status FROM usuarios WHERE nome = ? AND senha = ?",
                    (usuario, senha),
                )
                user = cursor.fetchone()
                conn.close()

                if user:
                    if user[5] == "Inativo":
                        st.error("⚠️ Este usuário está inativo.")
                    else:
                        st.session_state["logado"] = True
                        st.session_state["usuario"] = user[1]
                        st.session_state["perfil"] = user[2]
                        st.session_state["perm_ops"] = user[3]
                        st.session_state["perm_deps"] = user[4]
                        st.success("Acesso autorizado!")
                        st.rerun()
                else:
                    st.error("Credenciais inválidas.")
else:
    st.sidebar.title("Grupo Lima")
    st.sidebar.caption(f"Usuário: **{st.session_state['usuario']}** | Perfil: **{st.session_state['perfil']}**")

    unidade = st.sidebar.selectbox("Unidade / Operação", OPERACOES_DISPONIVEIS, key="sb_unidade_ativa")
    st.sidebar.divider()

    deps_disponiveis = DEPARTAMENTOS_DISPONIVEIS.copy()
    if st.session_state["perfil"] == "Master":
        deps_disponiveis.append("Acesso Master (Gestão de Usuários)")

    for d_name in deps_disponiveis:
        is_active = (st.session_state["nav_stack"][-1] == d_name)
        tipo_botao = "primary" if is_active else "secondary"
        if st.sidebar.button(d_name, key=f"sidebar_btn_{d_name}", use_container_width=True, type=tipo_botao):
            navigate_to(d_name)
            st.rerun()

    st.sidebar.divider()
    if st.sidebar.button("Sair do Sistema", use_container_width=True, key="btn_logout"):
        st.session_state["logado"] = False
        st.session_state["nav_stack"] = ["Visão Geral (Dashboard)"]
        st.rerun()

    c_back, c_title = st.columns([1, 8])
    if len(st.session_state["nav_stack"]) > 1:
        if c_back.button("⬅️ Voltar", key="btn_voltar_hist"):
            go_back()
            st.rerun()

    c_title.title(f"{st.session_state['nav_stack'][-1]}")
    st.caption(f"Operação ativa: **{unidade}**")

    dept_atual = st.session_state["nav_stack"][-1]

    if "Visão Geral" in dept_atual:
        st.subheader("Painel Geral de Desempenho Operacional")
        st.info("Utilize a barra lateral para navegar entre os módulos integrados.")

    elif "Ressuprimento" in dept_atual:
        sub_ress = st.tabs([
            "📁 Cadastros & Atualização de Bases",
            "📊 Gestão de Estoque",
            "🛒 Sugestão de Compra & Marcação por Dia",
            "📈 Gestão Ressuprimento (Cestas & Metas)",
        ])

        with sub_ress[3]:
            c_f1, c_f2 = st.columns(2)
            ano_sel = c_f1.number_input(
                "Ano de Análise:",
                min_value=2024,
                max_value=2030,
                value=datetime.now().year,
                key="ano_acompanhamento",
            )

            meses_nomes = [
                "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
            ]

            meses_selecionados = c_f2.multiselect(
                "Selecione os Meses de Análise (vazio = Ano Inteiro):",
                options=list(range(1, 13)),
                format_func=lambda x: meses_nomes[x - 1],
                default=[datetime.now().month],
                key="meses_acompanhamento",
            )

            mapa_op_sistema = {
                "Lima Rio Verde": ["Lima Rio Verde", "Lima - Rio Verde", "Rio Verde"],
                "Lima Barreiras": ["Lima Bahia", "Barreiras", "Lima Barreiras"],
                "Lima São Félix": ["Lima Bahia Samavi", "Samavi", "São Félix", "Lima São Félix"],
                "Bahia": ["Lima Barreiras", "Lima São Félix", "Barreiras", "Samavi", "São Félix", "Lima Bahia", "Lima Bahia Samavi"],
            }
            nombres_filtro = mapa_op_sistema.get(unidade, [unidade])
            nome_exibicao_op = "Bahia (Barreiras + São Félix)" if unidade == "Bahia" else unidade.replace("Lima ", "")

            conn = sqlite3.connect("puxada_ambev.db")
            placeholders_op = ",".join(["?"] * len(nombres_filtro))
            df_diario = pd.read_sql_query(
                f"SELECT * FROM gestao_ressuprimento_diario WHERE operacao IN ({placeholders_op}) AND strftime('%Y', data_registro)='{ano_sel}'",
                conn,
                params=nombres_filtro,
            )
            df_metas = pd.read_sql_query(
                f"SELECT * FROM metas_ressuprimento_mensal WHERE operacao IN ({placeholders_op}) AND ano={ano_sel}",
                conn,
                params=nombres_filtro,
            )
            conn.close()

            if not df_diario.empty:
                df_diario["data_dt"] = pd.to_datetime(
                    df_diario["data_registro"], errors="coerce"
                )

                if not meses_selecionados:
                    meses_selecionados = list(range(1, 13))

                df_diario_filtrado = df_diario[
                    df_diario["data_dt"].dt.month.isin(meses_selecionados)
                ]

                dias_preenchidos = df_diario_filtrado["data_dt"].dt.date.nunique()

                dias_no_ano_total = 0
                for m_s in meses_selecionados:
                    if m_s in [1, 3, 5, 7, 8, 10, 12]:
                        dias_no_ano_total += 31
                    elif m_s in [4, 6, 9, 11]:
                        dias_no_ano_total += 30
                    else:
                        dias_no_ano_total += (
                            29
                            if (
                                ano_sel % 4 == 0
                                and (ano_sel % 100 != 0 or ano_sel % 400 == 0)
                            )
                            else 28
                        )

                # Mapeamento correto da coluna volume_real_hl (Coluna C) para a coluna REAL
                df_res_mes = (
                    df_diario_filtrado.groupby("cesta")["volume_real_hl"]
                    .sum()
                    .reset_index()
                )

                cestas_map = {
                    "CATEGORIA_AGRUPADO - CERVEJA": "Cerveja",
                    "CATEGORIA_AGRUPADO - NAB": "Nab",
                    "CATEGORIA - MATCH": "Match",
                    "CATEGORIA_RETORNAVEL - CERVEJA RGB": "Cerveja RGB",
                    "REFRIGERANTE_REGULAR_NAB - ZERO": "Nab Zero",
                    "CERV_2 - Zero Alcool": "Cerveja Zero Alcool",
                    "SEGMENTO - HIGH END": "High End",
                }
                cestas_ordenadas = list(cestas_map.keys())

                df_comp = pd.merge(
                    pd.DataFrame({"cesta": cestas_ordenadas}),
                    df_res_mes,
                    on="cesta",
                    how="left",
                ).fillna(0)

                df_metas_filtradas = df_metas[
                    df_metas["mes"].isin(meses_selecionados)
                ]
                df_metas_grp = (
                    df_metas_filtradas.groupby("cesta")["meta_volume_hl"]
                    .sum()
                    .reset_index()
                )

                df_comp = pd.merge(
                    df_comp, df_metas_grp, on="cesta", how="left"
                ).fillna(0)

                fator_tend = (
                    (dias_no_ano_total / dias_preenchidos)
                    if dias_preenchidos > 0
                    else 1.0
                )

                df_comp["INDICADOR"] = df_comp["cesta"].map(cestas_map)
                df_comp["META"] = df_comp["meta_volume_hl"]
                df_comp["REAL"] = df_comp["volume_real_hl"]
                df_comp["TEND."] = df_comp["REAL"] * fator_tend

                df_comp["ATING. REAL"] = df_comp.apply(
                    lambda r: (r["REAL"] / r["META"] * 100)
                    if r["META"] > 0
                    else 0.0,
                    axis=1,
                )
                df_comp["ATING. TEND."] = df_comp.apply(
                    lambda r: (r["TEND."] / r["META"] * 100)
                    if r["META"] > 0
                    else 0.0,
                    axis=1,
                )

                df_comp["PENDÊNCIA PERÍODO"] = df_comp["REAL"] - df_comp["META"]

                df_cerveja_nab = df_comp[
                    df_comp["INDICADOR"].isin(["Cerveja", "Nab"])
                ]
                tot_meta = df_cerveja_nab["META"].sum()
                tot_real = df_cerveja_nab["REAL"].sum()
                tot_tend = df_cerveja_nab["TEND."].sum()
                tot_pend = tot_real - tot_meta

                tot_ating_real = (
                    (tot_real / tot_meta * 100) if tot_meta > 0 else 0.0
                )
                tot_ating_tend = (
                    (tot_tend / tot_meta * 100) if tot_meta > 0 else 0.0
                )

                df_total = pd.DataFrame([{
                    "INDICADOR": f"Total {nome_exibicao_op} (Cerveja + Nab)",
                    "META": tot_meta,
                    "REAL": tot_real,
                    "TEND.": tot_tend,
                    "ATING. REAL": tot_ating_real,
                    "ATING. TEND.": tot_ating_tend,
                    "PENDÊNCIA PERÍODO": tot_pend,
                }])

                meses_str = ", ".join([
                    meses_nomes[m - 1] for m in meses_selecionados
                ])
                st.markdown(
                    f"""
                    <div style="background-color: #0d2149; color: white; padding: 12px 20px; border-radius: 8px 8px 0 0; margin-bottom: 0px;">
                        <h3 style="margin:0; font-size: 20px; color: white;">🔵 {nome_exibicao_op}</h3>
                        <span style="font-size: 13px; color: #b0c4de;">Somatório de HL do Mês · Meses: {meses_str} · {dias_preenchidos} dia(s) computado(s)</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                df_final = pd.concat([
                    df_comp[[
                        "INDICADOR",
                        "META",
                        "REAL",
                        "TEND.",
                        "ATING. REAL",
                        "ATING. TEND.",
                        "PENDÊNCIA PERÍODO",
                    ]],
                    df_total,
                ], ignore_index=True)

                def format_atingimento_com_condicional(val):
                    try:
                        v = float(val)
                        s_val = f"{v:.1f}%".replace(".", ",")
                        if v >= 100.0:
                            return f"🟢 {s_val} (Batendo)"
                        else:
                            return f"🔴 {s_val} (Abaixo)"
                    except Exception:
                        return val

                df_view = df_final.copy()
                df_view["META"] = df_view["META"].apply(formatar_inteiro_br)
                df_view["REAL"] = df_view["REAL"].apply(formatar_inteiro_br)
                df_view["TEND."] = df_view["TEND."].apply(formatar_inteiro_br)
                df_view["ATING. REAL"] = df_view["ATING. REAL"].apply(
                    format_atingimento_com_condicional
                )
                df_view["ATING. TEND."] = df_view["ATING. TEND."].apply(
                    format_atingimento_com_condicional
                )
                df_view["PENDÊNCIA PERÍODO"] = df_view["PENDÊNCIA PERÍODO"].apply(
                    formatar_inteiro_br
                )

                st.dataframe(
                    df_view,
                    use_container_width=True,
                    height=(len(df_view) + 1) * 38 + 5,
                )

                st.markdown("##### 📥 Opções de Download do Relatório Mensal")
                render_botoes_download(df_final, f"Acompanhamento_Mensal_{unidade}")

            else:
                st.info(
                    f"ℹ️ Verifique se há dados diários cadastrados para **{nome_exibicao_op}** no ano de {ano_sel}."
                )

    elif "Financeiro" in dept_atual:
        st.subheader("💰 Gestão Financeira, Contas a Pagar, Vencimentos & Fluxo de Caixa Avançado")
        sub_fin = st.tabs([
            "📂 Relatório Diário",
            "💳 Contas a Pagar",
            "⏳ Gestão de Vencimentos",
            "🤖 Análise IA",
            "📊 Fluxo de Caixa Diário Avançado"
        ])

        with sub_fin[4]:
            st.markdown("### 📊 Fluxo de Caixa Diário Avançado (Próximos 15 Dias)")
            
            conn = sqlite3.connect("puxada_ambev.db")
            df_fc_base = pd.read_sql_query(f"SELECT * FROM financeiro_contas_pagar WHERE operacao = '{unidade}'", conn)
            conn.close()

            if not df_fc_base.empty:
                df_fc_base["data_venc_dt"] = pd.to_datetime(df_fc_base["data_vencimento"], errors="coerce").dt.date
                df_fc_base["Pendente"] = df_fc_base["comprometido"] - df_fc_base["realizado"]
                df_Ambev = df_fc_base[df_fc_base["nome_fornecedor"].str.upper().str.contains("AMBEV", na=False)]
                df_Sem_Ambev = df_fc_base[~df_fc_base["nome_fornecedor"].str.upper().str.contains("AMBEV", na=False)]
            else:
                df_Ambev = pd.DataFrame()
                df_Sem_Ambev = pd.DataFrame()

            st.markdown("##### ⚙️ Configurações Diárias e Saldo Banco Inicial")
            
            c_conf1, c_conf2 = st.columns([2, 1])
            saldo_inicial_input = c_conf1.number_input("💵 Saldo Banco Inicial (Dia Base):", value=50000.0, step=1000.0, format="%.2f", key="num_saldo_inicial_fc")
            
            if c_conf2.button("🔄 Atualizar com Contas a Pagar"):
                tot_ambev_total = df_Ambev["Pendente"].sum() if not df_Ambev.empty else 0.0
                st.success(f"Saldo sincronizado! Total Ambev pendente: R$ {formatar_br(tot_ambev_total)}")

            hoje_dt = datetime.now().date()
            dias_proj_lista = [hoje_dt + timedelta(days=i) for i in range(15)]

            cp_hoje_sem = df_Sem_Ambev[df_Sem_Ambev["data_venc_dt"] == hoje_dt]["Pendente"].sum() if not df_Sem_Ambev.empty else 0.0
            cp_hoje_ambev = df_Ambev[df_Ambev["data_venc_dt"] == hoje_dt]["Pendente"].sum() if not df_Ambev.empty else 0.0
            saldo_hoje_proj = saldo_inicial_input - cp_hoje_sem - cp_hoje_ambev

            st.markdown("---")
            st.markdown(f"""
                <div class="card-vibrante-verde" style="display: flex; justify-content: space-between; align-items: center;">
                    <div>
                        <h4 style="margin:0; color: white;">💰 Saldo Projetado para Hoje ({hoje_dt.strftime('%d/%m/%Y')})</h4>
                        <span style="font-size: 14px; color: #f1f5f9;">Calculado com base no Saldo Inicial e Contas Vencendo Hoje</span>
                    </div>
                    <div style="font-size: 24px; font-weight: bold; color: white;">
                        R$ {formatar_br(saldo_hoje_proj).replace('R$ ', '')}
                    </div>
                </div>
            """, unsafe_allow_html=True)

            st.divider()
            st.markdown("##### 📋 Tabela de Projeção Diária (Próximos 15 Dias)")

            hc1, hc2, hc3, hc4, hc5, hc6 = st.columns(6)
            hc1.markdown("**📅 Data**")
            hc2.markdown("**💵 Saldo Banco (R$)**")
            hc3.markdown("**🛒 Compra Ambev (R$)**")
            hc4.markdown("**📦 Contas Outros (R$)**")
            hc5.markdown("**📈 Previsão Entrada (R$)**")
            hc6.markdown("**💰 Saldo Projetado (R$)**")

            tabela_resultados_fc = []
            saldo_corr = saldo_inicial_input

            for d_item in dias_proj_lista:
                d_str = d_item.strftime("%Y-%m-%d")
                
                cp_outros_dia = df_Sem_Ambev[df_Sem_Ambev["data_venc_dt"] == d_item]["Pendente"].sum() if not df_Sem_Ambev.empty else 0.0
                cp_ambev_db = df_Ambev[df_Ambev["data_venc_dt"] == d_item]["Pendente"].sum() if not df_Ambev.empty else 0.0

                cols_d = st.columns(6)
                cols_d[0].markdown(f"**{d_item.strftime('%d/%m/%Y')}**")
                
                sb_atual = cols_d[1].number_input(f"Saldo #{d_str}", value=0.0, step=1000.0, format="%.2f", key=f"sb_{d_str}", label_visibility="collapsed")
                cp_ambev_man = cols_d[2].number_input(f"Ambev #{d_str}", value=float(cp_ambev_db), step=500.0, format="%.2f", key=f"ambev_{d_str}", label_visibility="collapsed")
                cp_outros_val = cols_d[3].number_input(f"Outros #{d_str}", value=float(cp_outros_dia), step=500.0, format="%.2f", key=f"outros_{d_str}", label_visibility="collapsed")
                prev_rec = cols_d[4].number_input(f"Rec #{d_str}", value=0.0, step=500.0, format="%.2f", key=f"rec_{d_str}", label_visibility="collapsed")

                if sb_atual > 0:
                    saldo_corr = sb_atual + prev_rec - cp_outros_val - cp_ambev_man
                else:
                    saldo_corr = saldo_corr + prev_rec - cp_outros_val - cp_ambev_man

                cols_d[5].markdown(f"**{formatar_br(saldo_corr)}**")

                tabela_resultados_fc.append({
                    "Data": d_item.strftime("%d/%m/%Y"),
                    "Saldo Banco": sb_atual,
                    "Compra Ambev (Contas a Pagar)": cp_ambev_man,
                    "Contas Outros Fornecedores": cp_outros_val,
                    "Previsão Entrada": prev_rec,
                    "Saldo Projetado do Dia": saldo_corr
                })

            df_res_fluxo = pd.DataFrame(tabela_resultados_fc)
            st.divider()
            render_botoes_download(df_res_fluxo, f"Fluxo_Caixa_Avancado_{unidade}")

    elif "Relatórios" in dept_atual:
        st.subheader("Base de Dados Completa para Download")
        tabela = st.selectbox("Escolha a tabela:", ["financeiro_contas_pagar", "cotacoes_frete"], key="sb_tabela_global_export")
        conn = sqlite3.connect("puxada_ambev.db")
        df = pd.read_sql_query(f"SELECT * FROM {tabela}", conn)
        conn.close()
        st.dataframe(df, use_container_width=True)
        render_botoes_download(df, f"Tabela_{tabela}")
